import io
import os
import numpy as np
import tensorflow as tf
from PIL import Image
from tensorflow.keras.applications.vgg16 import preprocess_input # Add this import
from fastapi.middleware.cors import CORSMiddleware
from huggingface_hub import hf_hub_download
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import StreamingResponse
from google import genai
from google.genai import types
from fpdf import FPDF
from markdown import markdown
from docx import Document
from bs4 import BeautifulSoup
import arabic_reshaper
from bidi.algorithm import get_display
from fpdf import FPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_markdown_to_word(document, md_text):
    """
    Converts Markdown to HTML and then injects it into a 
    python-docx document with basic formatting.
    """
    html = markdown(md_text)
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'li']):
        if element.name == 'h1':
            document.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            document.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            document.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            p = document.add_paragraph()
            # Handle bold/italic inside paragraphs
            for child in element.children:
                if child.name == 'strong' or child.name == 'b':
                    p.add_run(child.get_text()).bold = True
                elif child.name == 'em' or child.name == 'i':
                    p.add_run(child.get_text()).italic = True
                else:
                    p.add_run(str(child))
        elif element.name == 'ul':
            for li in element.find_all('li'):
                document.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                document.add_paragraph(li.get_text(), style='List Number')

GEMINI_API_KEY = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"  # Here replace with your actual API key
client = genai.Client(api_key=GEMINI_API_KEY)

app = FastAPI()

# Enable CORS for your HTML frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

model = None
# This model specifically classifies into these two categories
labels = ["Non-Cancerous", "Cancerous"]

print("--- Downloading Specific Keras Model File ---")
try:
    # 1. Download ONLY the .keras file from the repository
    model_path = hf_hub_download(
        repo_id="VRJBro/skin-cancer-detection", 
        filename="skin_cancer_model.keras"
    )
    
    # 2. Load it using the native Keras 3 loader
    model = tf.keras.models.load_model(model_path)
    print(f"✅ SUCCESS: Model loaded from {model_path}")
    
except Exception as e:
    print(f"❌ LOADING ERROR: {e}")


def fix_text(text):
    """
    Reshapes Arabic text to connect letters and reverses it for RTL display.
    """
    try:
        reshaped_text = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped_text)
        return bidi_text
    except:
        return text

class PDF(FPDF):
    def header(self):
        # NOTE: You must register the font in the setup first
        # We use the font name 'ArabicFont' defined below
        try:
            self.set_font("ArabicFont", "", 12) 
        except:
            self.set_font("Arial", "B", 12) # Fallback if font fails
            
        self.cell(0, 10, "AI Medical Analysis Report", border=False, ln=True, align="C")
        self.ln(5)

@app.get("/")
async def health():
    return {"status": "online", "model_loaded": model is not None}


@app.post("/analyze")
async def analyze(file: UploadFile = File(...)):
    """
    Step 1: Lightweight Classification
    Only runs the local VGG16 model. Returns JSON immediately.
    """
    if model is None:
        raise HTTPException(status_code=500, detail="Model not loaded.")

    try:
        contents = await file.read()
        img = Image.open(io.BytesIO(contents)).convert('RGB')
        img = img.resize((224, 224))
        
        # 1. Preprocessing (Your existing logic)
        img_array = np.array(img).astype(np.float32)
        img_array = np.expand_dims(img_array, axis=0)
        img_array = preprocess_input(img_array)

        # 2. Prediction
        predictions = model.predict(img_array)
        score = float(predictions[0][0])
        
        print(f"DEBUG: Raw model score is {score}")

        # 3. Label Logic
        if score > 0.5:
            label = "Cancerous"
            confidence = score
        else:
            label = "Non-Cancerous"
            confidence = 1 - score

        # Return JSON only - NO Gemini call here
        return {
            "prediction": label,
            "confidence": confidence
        }

    except Exception as e:
        print(f"Prediction Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate_report")
async def generate_report(
    file: UploadFile = File(...),
    prediction: str = Form(...),
    confidence: float = Form(...),
    language: str = Form(...),
    report_type: str = Form(...)
):
    """
    Step 2: Heavy Generation (On Demand)
    Calls Gemini API -> Generates PDF/Word -> Returns File
    """
    try:
        # 1. Read the image again for Gemini context
        contents = await file.read()

        # 2. Construct the Gemini Prompt (Moved from analyze)
        prompt = f"""
        ACT AS: An expert Dermatologist and Medical Consultant.
        CONTEXT: A patient has uploaded a dermoscopic image for skin lesion analysis. 
        AI MODEL RESULT: The local analysis model has classified the lesion as '{prediction}' with a confidence of {float(confidence)*100:.2f}%.

        TASK: Generate a comprehensive medical analysis report in {language}. 

        REPORT STRUCTURE (Use Markdown):
        1. ## Patient Report Summary
           - State the classification result clearly.
           - Explain what '{prediction}' generally means in simple terms.
        2. ## Clinical Observations
           - Based on the image analysis, describe what clinical features a specialist would look for (e.g., symmetry, borders, color variations).
        3. ## Recommendations & Next Steps
           - Provide clear actionable advice (e.g., "Monitor for changes using the ABCDE rule", "Consult a specialist for a biopsy", "Annual skin checks").
        4. ## Important Disclaimer
           - Include a standard medical disclaimer: This is an AI-generated report for educational purposes and NOT a final diagnosis.

        TONE: Professional, supportive, and clinical.
        FORMATTING: Use clear headings, bullet points, and bold text for emphasis.
        """

        # 3. Call Gemini API
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=[types.Part.from_bytes(data=contents, mime_type="image/jpeg"), prompt]
        )
        report_text = response.text

        # --- EXISTING PDF/WORD GENERATION LOGIC PRESERVED BELOW ---

        FONT_PATH = "fonts/Amiri-Regular.ttf" 

        if report_type == "pdf":
            pdf = PDF()
            
            # 1. REGISTER THE FONT (Crucial Step)
            # You need a .ttf file that supports Arabic.
            try:
                pdf.add_font('ArabicFont', '', FONT_PATH, uni=True)
                pdf.add_font('ArabicFont', 'B', FONT_PATH, uni=True)
            except FileNotFoundError:
                raise HTTPException(status_code=500, detail="Font file not found. Please add Amiri-Regular.ttf to your project.")

            pdf.add_page()
            
            # 2. Add Result Summary
            pdf.set_font("ArabicFont", '', 14)
            
            # Fix text direction for the Label
            display_label = fix_text(prediction.upper()) if language == "Arabic" else prediction.upper()

            color = (220, 38, 38) if prediction.lower() == "cancerous" else (22, 163, 74)
            pdf.set_text_color(*color)
            pdf.cell(0, 10, f"Analysis Result: {display_label}", ln=True, align='R' if language == "Arabic" else 'L')
            
            pdf.set_font("ArabicFont", '', 11)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(0, 10, f"Confidence Score: {float(confidence)*100:.2f}%", ln=True, align='R' if language == "Arabic" else 'L')
            
            pdf.ln(5)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)

            # 3. Handle The Main Report Text
            # HTML conversion + Arabic is very buggy in FPDF.
            # It is safer to write it as standard MultiCell text for Arabic to ensure characters connect.
            
            pdf.set_font("ArabicFont", '', 11)
            
            if language == "Arabic":
                # For Arabic, we clean markdown syntax (*, #) because simple FPDF text doesn't render MD
                clean_text = report_text.replace("**", "").replace("##", "").replace("- ", "• ")
                
                # Fix Directionality
                fixed_text = fix_text(clean_text)
                
                # Write aligned to Right
                pdf.multi_cell(0, 10, fixed_text, align='R')
            else:
                # Non-Arabic languages can use the HTML renderer
                html_content = markdown(report_text)
                pdf.write_html(html_content)

            pdf_bytes = pdf.output()
                
            return StreamingResponse(
                io.BytesIO(pdf_bytes), 
                media_type="application/pdf",
                headers={
                    "Content-Disposition": "attachment; filename=Report.pdf",
                    "Cache-Control": "no-cache"
                }
            )    

        # --- WORD DOCUMENT SUPPORT ---
        if report_type == "word":
            doc = Document()
            
            # Add Header
            title = doc.add_heading('Medical Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if language == "Arabic" else WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(f"FINAL RESULT: {prediction.upper()}")
            run.bold = True
            run.font.size = 140000 
            
            # Enable RTL for the run if needed
            if language == "Arabic":
                p.paragraph_format.bidi = True
            
            doc.add_paragraph(f"Confidence Score: {float(confidence)*100:.2f}%")
            doc.add_section()

            # Content
            # python-docx handles Unicode automatically, but we need to set RTL direction
            
            # Simple markdown stripper for Word (or use a library like 'markdown2docx')
            lines = report_text.split('\n')
            for line in lines:
                if line.strip():
                    para = doc.add_paragraph(line.replace("**", "").replace("##", ""))
                    if language == "Arabic":
                        para.paragraph_format.bidi = True # Sets RTL direction
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            out = io.BytesIO()
            doc.save(out)
            out.seek(0)

            return StreamingResponse(
                out, 
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=Analysis_Report.docx"}
            )

    except Exception as e:
        print(f"Report Generation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))
